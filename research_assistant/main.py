import os
import time
import io
import logging
from datetime import datetime
from dotenv import load_dotenv

# Fix SSL certificate loading issue with litellm
os.environ['REQUESTS_CA_BUNDLE'] = ''
os.environ['CURL_CA_BUNDLE'] = ''
os.environ['SSL_CERT_FILE'] = ''

# Load environment variables FIRST
load_dotenv()

from fastapi import FastAPI, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Any, Dict, Optional
import aiohttp
from bs4 import BeautifulSoup
import json
import re

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# File parsing libraries
import PyPDF2
from docx import Document as DocxDocument
import openpyxl
import csv

from google.adk import Runner
from google.adk.sessions import InMemorySessionService
from google.genai import types
from research_assistant import root_agent

# Web search tools (with DuckDuckGo fallback)
from tools.web_search import web_search, format_search_results_for_agent

# Check API key
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key or api_key == "your_api_key_here":
    logger.warning("⚠️  GOOGLE_API_KEY not configured!")
    print("\n" + "="*60)
    print("⚠️  WARNING: GOOGLE_API_KEY not configured!")
    print("="*60)
    print("\nTo get your API key:")
    print("1. Go to: https://aistudio.google.com/app/apikey")
    print("2. Create a new API key")
    print("3. Add it to the .env file")
    print("\n" + "="*60 + "\n")
else:
    logger.info("✅ API Key configured successfully")

# Initialize FastAPI
app = FastAPI(
    title="Research Assistant - Google ADK Sequential Pipeline",
    description="Multi-agent research assistant built with Google ADK SequentialAgent pattern",
    version="2.0.0"
)

# Mount static files
static_dir = os.path.join(os.path.dirname(__file__), "static")
os.makedirs(static_dir, exist_ok=True)
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Favicon route
@app.get("/favicon.ico")
async def favicon():
    return FileResponse(os.path.join(static_dir, "favicon.svg"), media_type="image/svg+xml")

# Session service for ADK
session_service = InMemorySessionService()
logger.info("✅ Session service initialized")


# ============================================
# File Parsing Functions
# ============================================
def extract_text_from_pdf(file_content: bytes) -> tuple[str, int]:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                # Clean any problematic characters
                clean_text = text.encode('utf-8', errors='ignore').decode('utf-8')
                text_parts.append(clean_text)
        return "\n\n".join(text_parts), len(pdf_reader.pages)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> tuple[str, int]:
    """Extract text from Word document"""
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                clean_text = para.text.encode('utf-8', errors='ignore').decode('utf-8')
                text_parts.append(clean_text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    clean_row = " | ".join(row_text)
                    text_parts.append(clean_row.encode('utf-8', errors='ignore').decode('utf-8'))
        
        return "\n\n".join(text_parts), 1
    except Exception as e:
        raise ValueError(f"Failed to parse Word document: {str(e)}")


def extract_text_from_excel(file_content: bytes) -> tuple[str, int]:
    """Extract text from Excel file"""
    workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
    text_parts = []
    sheet_count = 0
    
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        sheet_count += 1
        text_parts.append(f"=== Sheet: {sheet_name} ===")
        
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            if any(v.strip() for v in row_values):
                text_parts.append(" | ".join(row_values))
    
    return "\n".join(text_parts), sheet_count


# Request/Response Models
class ResearchRequest(BaseModel):
    question: str
    document: str
    enable_web_search: bool = False


class WebSearchRequest(BaseModel):
    query: str
    num_results: int = 5


class UrlFetchRequest(BaseModel):
    url: str


class UrlFetchResponse(BaseModel):
    success: bool
    text: Optional[str] = None
    title: Optional[str] = None
    error: Optional[str] = None


class WebSearchResponse(BaseModel):
    success: bool
    results: Optional[List[Dict[str, Any]]] = None
    query: Optional[str] = None
    error: Optional[str] = None


class ResearchResponse(BaseModel):
    success: bool
    steps: List[Dict[str, Any]]
    final_output: Optional[dict] = None
    error: Optional[str] = None
    web_search_used: bool = False


class FileUploadResponse(BaseModel):
    success: bool
    text: Optional[str] = None
    pages: Optional[int] = None
    file_type: Optional[str] = None
    error: Optional[str] = None


# Routes
@app.get("/")
async def root():
    """Serve the main HTML page"""
    logger.info("📄 Serving main page")
    return FileResponse(os.path.join(static_dir, "index.html"))


@app.post("/api/fetch-url", response_model=UrlFetchResponse)
async def fetch_url_content(request: UrlFetchRequest):
    """
    Fetch and extract content from a URL
    """
    logger.info(f"🔗 Fetching URL: {request.url}")
    
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        async with aiohttp.ClientSession() as session:
            async with session.get(request.url, headers=headers, timeout=30) as response:
                if response.status != 200:
                    logger.error(f"❌ URL fetch failed: HTTP {response.status}")
                    return UrlFetchResponse(
                        success=False,
                        error=f"Failed to fetch URL: HTTP {response.status}"
                    )
                
                html = await response.text()
                
                # Parse HTML and extract text
                soup = BeautifulSoup(html, 'html.parser')
                
                # Get title
                title = soup.title.string if soup.title else "Unknown"
                
                # Remove script and style elements
                for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                    element.decompose()
                
                # Get main content (try common content containers)
                main_content = soup.find('main') or soup.find('article') or soup.find('div', {'class': re.compile(r'content|article|post|entry')}) or soup.body
                
                if main_content:
                    text = main_content.get_text(separator='\n', strip=True)
                else:
                    text = soup.get_text(separator='\n', strip=True)
                
                # Clean up text
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                clean_text = '\n'.join(lines)
                
                logger.info(f"✅ URL fetched successfully: {len(clean_text)} characters from '{title}'")
                
                return UrlFetchResponse(
                    success=True,
                    text=clean_text[:50000],  # Limit to 50k chars
                    title=title
                )
                
    except aiohttp.ClientError as e:
        logger.error(f"❌ Network error: {str(e)}")
        return UrlFetchResponse(success=False, error=f"Network error: {str(e)}")
    except Exception as e:
        logger.error(f"❌ Error fetching URL: {str(e)}")
        return UrlFetchResponse(success=False, error=f"Error: {str(e)}")


@app.post("/api/upload", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    Upload and parse ANY document type
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx, .doc)
    - Excel (.xlsx, .xls)
    - PowerPoint (.pptx)
    - Text (.txt, .md, .rst)
    - CSV (.csv)
    - JSON (.json)
    - XML (.xml)
    - HTML (.html, .htm)
    - Code files (.py, .js, .java, .cpp, etc.)
    - And more...
    """
    logger.info(f"📁 Processing file upload: {file.filename}")
    
    try:
        # Read file content
        content = await file.read()
        filename = file.filename.lower()
        file_ext = os.path.splitext(filename)[1]
        
        logger.info(f"📄 File type detected: {file_ext}")
        
        text = ""
        pages = 1
        file_type = "unknown"
        
        # PDF files
        if file_ext == '.pdf':
            text, pages = extract_text_from_pdf(content)
            file_type = "PDF"
            
        # Word documents
        elif file_ext in ['.docx', '.doc']:
            text, pages = extract_text_from_docx(content)
            file_type = "Word Document"
            
        # Excel files
        elif file_ext in ['.xlsx', '.xls']:
            text, pages = extract_text_from_excel(content)
            file_type = "Excel Spreadsheet"
            
        # Text-based files (read directly)
        elif file_ext in ['.txt', '.md', '.rst', '.log', '.ini', '.cfg', '.conf']:
            text = content.decode('utf-8', errors='ignore')
            file_type = "Text File"
            
        # CSV files
        elif file_ext == '.csv':
            text_io = io.StringIO(content.decode('utf-8', errors='ignore'))
            reader = csv.reader(text_io)
            rows = [" | ".join(row) for row in reader]
            text = "\n".join(rows)
            file_type = "CSV"
            
        # JSON files
        elif file_ext == '.json':
            try:
                json_data = json.loads(content.decode('utf-8', errors='ignore'))
                text = json.dumps(json_data, indent=2)
                file_type = "JSON"
            except:
                text = content.decode('utf-8', errors='ignore')
                file_type = "JSON (raw)"
                
        # XML files
        elif file_ext == '.xml':
            text = content.decode('utf-8', errors='ignore')
            # Try to prettify
            try:
                soup = BeautifulSoup(text, 'xml')
                text = soup.prettify()
            except:
                pass
            file_type = "XML"
            
        # HTML files
        elif file_ext in ['.html', '.htm']:
            html_content = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_content, 'html.parser')
            for element in soup(['script', 'style']):
                element.decompose()
            text = soup.get_text(separator='\n', strip=True)
            file_type = "HTML"
            
        # Code files
        elif file_ext in ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.h', '.cs', '.go', '.rb', '.php', '.swift', '.kt', '.rs', '.sql', '.sh', '.bat', '.ps1', '.yaml', '.yml', '.toml']:
            text = content.decode('utf-8', errors='ignore')
            file_type = f"Code ({file_ext})"
            
        # Try to read as text for unknown formats
        else:
            try:
                text = content.decode('utf-8', errors='ignore')
                if text and len(text.strip()) > 0:
                    file_type = f"Text ({file_ext})"
                else:
                    # Binary file - can't process
                    logger.warning(f"⚠️ Could not extract text from {file_ext} file")
                    return FileUploadResponse(
                        success=False,
                        error=f"Cannot extract text from {file_ext} files. Please try a different format."
                    )
            except:
                return FileUploadResponse(
                    success=False,
                    error=f"Unsupported file format: {file_ext}"
                )
        
        # Check if text was extracted
        if not text or not text.strip():
            logger.warning("⚠️ No text extracted from file")
            return FileUploadResponse(
                success=False,
                error="Could not extract text from the file. The file may be empty or corrupted."
            )
        
        logger.info(f"✅ File processed: {len(text)} characters extracted ({file_type})")
        
        return FileUploadResponse(
            success=True,
            text=text,
            pages=pages,
            file_type=file_type
        )
        
    except UnicodeDecodeError as e:
        return FileUploadResponse(
            success=False,
            error="Could not decode file. Please ensure the file is not corrupted."
        )
    except Exception as e:
        # Safely convert error to string, handling any encoding issues
        error_msg = str(e).encode('utf-8', errors='ignore').decode('utf-8')
        return FileUploadResponse(
            success=False,
            error=f"Error processing file: {error_msg}"
        )


@app.post("/api/research")
async def process_research(request: ResearchRequest):
    """
    Process research question through ADK Sequential Pipeline
    
    Uses Google ADK SequentialAgent pattern:
    1. IntentDetectorAgent → Detect analysis mode
    2. DataProcessorAgent → Process document
    3. AnalyzerAgent → Search/Compare
    4. ExtractorAgent → Extract citations/scores
    5. ReportGeneratorAgent → Generate final report
    """
    logger.info("="*60)
    logger.info("🚀 NEW RESEARCH REQUEST")
    logger.info("="*60)
    logger.info(f"📝 Question: {request.question[:100]}...")
    logger.info(f"📄 Document: {len(request.document)} characters")
    logger.info(f"🌐 Web Search: {'Enabled' if request.enable_web_search else 'Disabled'}")
    
    try:
        # Check API key
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key or api_key == "your_api_key_here":
            logger.error("❌ API Key not configured")
            return ResearchResponse(
                success=False,
                steps=[],
                final_output=None,
                error="GOOGLE_API_KEY not configured. Please add your API key to the .env file. Get it from: https://aistudio.google.com/app/apikey"
            )
        
        # Validate inputs
        if not request.question.strip():
            logger.warning("⚠️ Empty question provided")
            return ResearchResponse(
                success=False,
                steps=[],
                error="Research question is required"
            )
        
        if not request.document.strip():
            logger.warning("⚠️ Empty document provided")
            return ResearchResponse(
                success=False,
                steps=[],
                error="Document content is required"
            )
        
        # Generate unique session ID for this request
        import uuid
        session_id = str(uuid.uuid4())
        user_id = "user_1"
        app_name = "research_assistant"
        
        logger.info(f"🔑 Session ID: {session_id[:8]}...")
        
        # Pre-process document to add location markers
        logger.info("📍 Pre-processing document for location tracking...")
        from agents.document_processor_agent import DocumentProcessorAgent
        from models.data_models import AgentInput
        
        doc_processor = DocumentProcessorAgent()
        input_data = AgentInput(data={
            "document": request.document,
            "question": request.question
        })
        processed_output = await doc_processor.execute(input_data)
        
        if processed_output.status == "error":
            logger.error(f"❌ Document processing failed: {processed_output.errors}")
            return ResearchResponse(
                success=False,
                steps=[],
                error=f"Document processing failed: {'; '.join(processed_output.errors)}"
            )
        
        # Get the indexed document with location markers
        indexed_document = processed_output.data.get("indexed_document", request.document)
        logger.info(f"✅ Document indexed with {processed_output.data.get('paragraph_count', 0)} paragraphs")
        
        # Create session for this request
        session = await session_service.create_session(
            app_name=app_name,
            user_id=user_id,
            session_id=session_id
        )
        logger.info("✅ Session created")
        
        # Create runner for the ADK sequential agent
        runner = Runner(
            agent=root_agent,
            app_name=app_name,
            session_service=session_service
        )
        
        # Prepare input message for the pipeline with INDEXED document
        user_message = f"""RESEARCH QUESTION: {request.question}

DOCUMENT TO ANALYZE (with location markers [P#:L#-#]):
{indexed_document}

⚠️ CRITICAL INSTRUCTIONS FOR PROVIDING ACTIONABLE FIXES:
========================================================

1. Analyze this document and answer the research question thoroughly.

2. **FOR EVERY ISSUE/IMPROVEMENT YOU IDENTIFY, YOU MUST PROVIDE:**
   
   📍 Location: [P#:L#-#]
   
   ❌ Current (Incorrect/Needs Improvement):
   ```
   [exact quote from document]
   ```
   
   ✅ Fixed/Improved Version (COPY-PASTE READY):
   ```
   [corrected version that user can directly copy]
   ```
   
   📝 Explanation: [Why this fix works in simple terms]
   🔴 Severity: [High/Medium/Low]

3. **Examples of fixes to provide:**
   - Code bugs → Working corrected code
   - Security issues → Secure implementation
   - Grammar errors → Corrected text
   - Performance problems → Optimized solution
   - Logic errors → Fixed logic
   - Style issues → Improved formatting

4. EXPLAIN EVERYTHING IN SIMPLE, LAYMAN TERMS that anyone can understand.
5. Avoid jargon - if you must use technical terms, explain them simply.
6. Use examples and analogies to make complex concepts clear.
7. If you cannot find relevant information, respond with "Couldn't fetch data from the given records".

⚠️ REMEMBER: Always provide the FIXED/CORRECTED version, not just descriptions!"""
        
        logger.info("📤 Sending to ADK Pipeline with location tracking...")
        
        # Create proper Content object for ADK (required format)
        new_message = types.Content(
            role="user",
            parts=[types.Part(text=user_message)]
        )
        
        # Track execution
        steps = []
        final_response = None
        start_time = time.time()
        
        logger.info("🔄 Starting Sequential Agent Pipeline...")
        
        # Run the ADK Sequential Pipeline (5 agents)
        async for event in runner.run_async(
            user_id=user_id,
            session_id=session_id,
            new_message=new_message
        ):
            step_time = time.time() - start_time
            
            # Capture agent events
            if hasattr(event, 'author') and event.author:
                agent_name = event.author
                event_content = ""
                
                if hasattr(event, 'content') and event.content:
                    if hasattr(event.content, 'parts'):
                        event_content = "".join([p.text for p in event.content.parts if hasattr(p, 'text')])
                    else:
                        event_content = str(event.content)
                
                # Track all 5 universal agents
                tracked_agents = [
                    'IntentDetectorAgent',
                    'DataProcessorAgent', 
                    'AnalyzerAgent',
                    'ExtractorAgent',
                    'ReportGeneratorAgent'
                ]
                if agent_name in tracked_agents and event_content:
                    logger.info(f"🤖 {agent_name} completed ({step_time:.2f}s)")
                    preview = event_content[:300] + "..." if len(event_content) > 300 else event_content
                    
                    steps.append({
                        "agent": agent_name,
                        "status": "success",
                        "output_preview": preview,
                        "processing_time": round(step_time, 3)
                    })
                    
                    start_time = time.time()
                
                # Capture final response (from ReportGeneratorAgent)
                if agent_name == 'ReportGeneratorAgent' and event_content:
                    final_response = event_content
                    logger.info("✅ Final report generated")
        
        # Parse final output
        if final_response:
            if "couldn't fetch data" in final_response.lower():
                logger.warning("⚠️ No relevant data found in document")
                final_output = {
                    "status": "no_data",
                    "answer": "Couldn't fetch data from the given records",
                    "explanation": "The document does not contain information relevant to your analysis. Try providing more relevant data or rephrasing your query."
                }
            else:
                logger.info(f"✅ Analysis complete: {len(final_response)} characters")
                # Parse the full response and format for frontend
                final_output = {
                    "status": "success",
                    "answer": final_response,
                    "citations": []
                }
        else:
            logger.error("❌ No response generated from pipeline")
            final_output = {
                "status": "no_data",
                "answer": "No response generated",
                "explanation": "The pipeline did not produce a response. Please try again."
            }
        
        total_time = sum([s.get('processing_time', 0) for s in steps])
        logger.info(f"⏱️ Total processing time: {total_time:.2f}s")
        logger.info("="*60)
        
        return ResearchResponse(
            success=True,
            steps=steps,
            final_output=final_output
        )
        
    except Exception as e:
        import traceback
        logger.error(f"❌ Pipeline error: {str(e)}")
        logger.error(traceback.format_exc())
        return ResearchResponse(
            success=False,
            steps=[],
            final_output=None,
            error=f"{str(e)}"
        )


@app.post("/api/search")
async def perform_web_search(request: WebSearchRequest):
    """
    Perform web search (Google with DuckDuckGo fallback)
    
    Uses Google Custom Search API if configured, otherwise DuckDuckGo (free)
    """
    try:
        if not request.query.strip():
            return WebSearchResponse(
                success=False,
                results=[],
                error="Search query is required"
            )
        
        # Perform web search (tries Google, falls back to DuckDuckGo)
        results = await web_search(
            query=request.query,
            num_results=request.num_results
        )
        
        if results.get("success"):
            return WebSearchResponse(
                success=True,
                results=results.get("results", []),
                query=request.query
            )
        else:
            return WebSearchResponse(
                success=False,
                results=[],
                error=results.get("error", "Search failed")
            )
        
    except Exception as e:
        return WebSearchResponse(
            success=False,
            results=[],
            error=str(e)
        )


@app.post("/api/research/enhanced")
async def process_research_enhanced(request: ResearchRequest):
    """
    Enhanced research with automatic comprehensive analysis
    
    When question is [AUTO-COMPREHENSIVE-ANALYSIS]:
    1. Automatically performs Summary + Literature Review + Competitive Analysis
    2. Web search is always enabled for comprehensive context
    3. Results are formatted in easy-to-understand layman terms
    """
    logger.info("="*60)
    logger.info("🚀 NEW ENHANCED RESEARCH REQUEST")
    logger.info("="*60)
    
    is_auto_analysis = request.question == "[AUTO-COMPREHENSIVE-ANALYSIS]"
    
    if is_auto_analysis:
        logger.info("🤖 AUTO-COMPREHENSIVE-ANALYSIS MODE")
        logger.info("📋 Will perform: Summary → Literature Review → Competitive Analysis")
    else:
        logger.info(f"📝 Question: {request.question[:100]}...")
    
    logger.info(f"📄 Document: {len(request.document)} characters")
    logger.info(f"🌐 Web Search: ENABLED")
    
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key or api_key == "your_api_key_here":
            logger.error("❌ API Key not configured")
            return ResearchResponse(
                success=False,
                steps=[],
                final_output=None,
                error="GOOGLE_API_KEY not configured"
            )
        
        web_context = ""
        web_search_used = False
        
        # Always perform web search for comprehensive analysis
        logger.info("🔍 Performing web search...")
        try:
            # Extract key topics from document for search
            doc_preview = request.document[:500]
            search_query = request.question if not is_auto_analysis else f"analyze {doc_preview[:200]}"
            
            search_results = await web_search(
                query=search_query,
                num_results=5
            )
            
            if search_results.get("success"):
                web_context = format_search_results_for_agent(search_results)
                web_search_used = True
                logger.info(f"✅ Web search successful via {search_results.get('source', 'Unknown')}")
        except Exception as e:
            logger.error(f"❌ Web search failed: {e}")
        
        # Generate unique session ID
        import uuid
        session_id = str(uuid.uuid4())
        user_id = "user_1"
        app_name = "research_assistant"
        
        logger.info(f"🔑 Session ID: {session_id[:8]}...")
        
        session = await session_service.create_session(
            app_name=app_name,
            user_id=user_id,
            session_id=session_id
        )
        logger.info("✅ Session created")
        
        runner = Runner(
            agent=root_agent,
            app_name=app_name,
            session_service=session_service
        )
        
        # Build comprehensive analysis message
        if is_auto_analysis:
            user_message = f"""⚠️ MANDATORY COMPREHENSIVE ANALYSIS - ALL SECTIONS REQUIRED ⚠️

CONTENT TO ANALYZE:
{request.document}

{f"WEB SEARCH CONTEXT:{chr(10)}{web_context}" if web_context else ""}

═══════════════════════════════════════════════════════════════
⚠️ YOU MUST PROVIDE ALL FOUR SECTIONS BELOW. DO NOT SKIP ANY! ⚠️
═══════════════════════════════════════════════════════════════

📝 SUMMARY (MANDATORY - DO NOT SKIP)
=====================================
Create a clear, comprehensive summary of this content. Include:
- Main topic and purpose
- Key points and findings  
- Important facts and figures
- Overall conclusions

📚 LITERATURE REVIEW (MANDATORY - DO NOT SKIP)
==============================================
Perform a literature review analysis:
1. GATHER: Identify all key themes, topics, and subjects covered
2. CATEGORIZE: Group related information into logical categories
3. ANALYZE: Examine the depth, quality, and significance of each topic
4. SYNTHESIZE: Connect the dots - how do different parts relate to each other?
- Identify any gaps or missing information
- Note any contrasting viewpoints

🏆 COMPETITIVE ANALYSIS (MANDATORY - DO NOT SKIP)
=================================================
Perform competitive/comparative analysis:
1. GATHER: Identify all entities, products, companies, methods, or options mentioned
2. COMPARE: List similarities and differences between them
3. SCORE: Rate each on relevant criteria (quality, features, value, etc.)
4. REPORT: Provide rankings and recommendations

**YOU MUST FORMAT COMPARISONS AS A MARKDOWN TABLE:**
| Feature | Option A | Option B | Option C |
|---------|----------|----------|----------|
| Feature 1 | Value | Value | Value |
| Quality | ⭐⭐⭐⭐ | ⭐⭐⭐ | ⭐⭐⭐⭐⭐ |
| Score | 8/10 | 6/10 | 9/10 |

If no direct competitors exist, compare:
- Different approaches/methods mentioned
- Pros vs Cons of the main topic
- Before vs After scenarios
- Different use cases or applications

� FIXES & IMPROVEMENTS (MANDATORY - DO NOT SKIP)
=================================================
**For every issue, error, or improvement opportunity found, provide:**

📍 Location: [P#:L#-#]

❌ Current (Incorrect/Needs Improvement):
```
[exact quote from document]
```

✅ Fixed/Improved Version (COPY-PASTE READY):
```
[corrected version that user can directly copy]
```

📝 Explanation: [Why this fix works in simple terms]
🔴 Severity: [High/Medium/Low]

**Types of fixes to identify:**
- Code bugs, security vulnerabilities, performance issues
- Grammar errors, spelling mistakes, punctuation problems
- Logic errors, inconsistencies, unclear statements
- Style improvements, formatting issues

If no issues found, state: "✅ No critical issues identified."

💡 KEY TAKEAWAYS (MANDATORY - DO NOT SKIP)
==========================================
List 3-5 actionable insights or recommendations.

═══════════════════════════════════════════════════════════════
CRITICAL INSTRUCTIONS:
═══════════════════════════════════════════════════════════════
1. ⭐ YOU MUST INCLUDE ALL 5 SECTIONS: Summary, Literature Review, Competitive Analysis, Fixes & Improvements, Key Takeaways
2. EXPLAIN EVERYTHING IN SIMPLE, LAYMAN TERMS that anyone can understand
3. Avoid jargon - if you must use technical terms, explain them in parentheses
4. Use bullet points, numbered lists, and clear headings for readability
5. ALWAYS include a comparison table in the Competitive Analysis section
6. **ALWAYS provide copy-paste ready fixes with before/after code/text, not just descriptions**
7. Make the analysis practical and actionable

FORMAT YOUR RESPONSE EXACTLY LIKE THIS:
📝 SUMMARY
[content]

📚 LITERATURE REVIEW
[content]

🏆 COMPETITIVE ANALYSIS
[content with table]

🔧 FIXES & IMPROVEMENTS
[fixes with before/after code/text]

💡 KEY TAKEAWAYS
[content]"""
        else:
            # Regular question-based analysis
            user_message = f"""RESEARCH QUESTION: {request.question}

DOCUMENT TO ANALYZE:
{request.document}"""
            
            if web_context:
                user_message += f"""

WEB SEARCH RESULTS (Additional Context):
{web_context}

IMPORTANT INSTRUCTIONS:
1. Analyze the document AND incorporate relevant insights from web search results.
2. EXPLAIN EVERYTHING IN SIMPLE, LAYMAN TERMS that anyone can understand.
3. Avoid jargon - if you must use technical terms, explain them simply.
4. Use examples and analogies to make complex concepts clear.
5. Provide a comprehensive, easy-to-understand analysis."""
            else:
                user_message += """

IMPORTANT INSTRUCTIONS:
1. Analyze this document and answer the research question thoroughly.
2. EXPLAIN EVERYTHING IN SIMPLE, LAYMAN TERMS that anyone can understand.
3. Avoid jargon - if you must use technical terms, explain them simply.
4. Use examples and analogies to make complex concepts clear."""
        
        logger.info("📤 Sending to ADK Pipeline...")
        
        new_message = types.Content(
            role="user",
            parts=[types.Part(text=user_message)]
        )
        
        steps = []
        final_response = None
        start_time = time.time()
        
        logger.info("🔄 Starting Sequential Agent Pipeline...")
        
        if web_search_used:
            logger.info("🌐 Adding web search results to pipeline")
            steps.append({
                "agent": "WebSearchAgent",
                "status": "success",
                "output_preview": f"Found web results for context enhancement",
                "processing_time": 0.5
            })
        
        async for event in runner.run_async(
            user_id=user_id,
            session_id=session_id,
            new_message=new_message
        ):
            step_time = time.time() - start_time
            
            if hasattr(event, 'author') and event.author:
                agent_name = event.author
                event_content = ""
                
                if hasattr(event, 'content') and event.content:
                    if hasattr(event.content, 'parts'):
                        event_content = "".join([p.text for p in event.content.parts if hasattr(p, 'text')])
                    else:
                        event_content = str(event.content)
                
                tracked_agents = [
                    'IntentDetectorAgent',
                    'DataProcessorAgent', 
                    'AnalyzerAgent',
                    'ExtractorAgent',
                    'ReportGeneratorAgent'
                ]
                if agent_name in tracked_agents and event_content:
                    preview = event_content[:300] + "..." if len(event_content) > 300 else event_content
                    
                    logger.info(f"🤖 {agent_name} completed ({step_time:.2f}s)")
                    steps.append({
                        "agent": agent_name,
                        "status": "success",
                        "output_preview": preview,
                        "processing_time": round(step_time, 3)
                    })
                    
                    start_time = time.time()
                
                if agent_name == 'ReportGeneratorAgent' and event_content:
                    final_response = event_content
                    logger.info("✅ Final report generated")
        
        if final_response:
            if "couldn't fetch data" in final_response.lower():
                logger.warning("⚠️ No relevant data found")
                final_output = {
                    "status": "no_data",
                    "answer": "Couldn't fetch data from the given records",
                    "explanation": "The document does not contain relevant information."
                }
            else:
                logger.info(f"✅ Analysis complete: {len(final_response)} characters")
                final_output = {
                    "status": "success",
                    "answer": final_response,
                    "citations": [],
                    "web_enhanced": web_search_used
                }
        else:
            logger.error("❌ No response generated")
            final_output = {
                "status": "no_data",
                "answer": "No response generated",
                "explanation": "Please try again."
            }
        
        total_time = sum([s.get('processing_time', 0) for s in steps])
        logger.info(f"⏱️ Total processing time: {total_time:.2f}s")
        logger.info("="*60)
        
        return ResearchResponse(
            success=True,
            steps=steps,
            final_output=final_output,
            web_search_used=web_search_used
        )
        
    except Exception as e:
        import traceback
        logger.error(f"❌ Pipeline error: {str(e)}")
        logger.error(traceback.format_exc())
        return ResearchResponse(
            success=False,
            steps=[],
            final_output=None,
            error=str(e)
        )


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    logger.info("💓 Health check requested")
    api_key = os.getenv("GOOGLE_API_KEY")
    search_cx = os.getenv("GOOGLE_SEARCH_CX")
    return {
        "status": "healthy",
        "service": "Universal Research Assistant",
        "framework": "Google ADK",
        "pattern": "SequentialAgent",
        "api_key_configured": bool(api_key and api_key != "your_api_key_here"),
        "web_search_configured": bool(search_cx),
        "agents": [
            "IntentDetectorAgent",
            "DataProcessorAgent",
            "AnalyzerAgent",
            "ExtractorAgent", 
            "ReportGeneratorAgent"
        ],
        "modes": ["Research Assistant", "Literature Review", "Competitive Analysis"]
    }


# Run server
if __name__ == "__main__":
    import uvicorn
    
    print("\n" + "="*60)
    print("🚀 Research Assistant - Google ADK Sequential Pipeline")
    print("="*60)
    print("\n📍 Open your browser: http://localhost:8000")
    print("\n🤖 Framework: Google ADK (Agent Development Kit)")
    print("📋 Pattern: SequentialAgent")
    print("⚡ LLM: Google Gemini")
    print("\n🔗 Pipeline Flow:")
    print("   1. DocumentProcessorAgent → Parse document")
    print("   2. SearchAgent → Find relevant info")
    print("   3. ExtractionAgent → Extract citations")
    print("   4. SummaryAgent → Generate response")
    print("\n" + "="*60 + "\n")
    
    uvicorn.run(app, host="127.0.0.1", port=8000)
